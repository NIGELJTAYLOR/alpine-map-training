"""
restyle_workbook.py
====================
Re-style the Alpine Map Training Companion Manual to the Glacier Lab design system.

Run from the directory containing this script and the source .docx:

    pip install python-docx
    python restyle_workbook.py

By default reads:   Alpine_Map_Training_Companion_Manual_v11.docx
Writes:             Alpine_Map_Training_Companion_Manual_v11_glacier.docx

This script makes ONE pass:
  1. Defines / redefines paragraph styles to Glacier Lab values
  2. Walks every paragraph; re-tags by text pattern + existing style
  3. Strips inline run-property overrides that force Calibri on everything
  4. (Optional) inserts a cover page at the top
  5. (Optional) appends an Answer Key stub at the back

Iterate by tweaking the constants in the TOKENS section at the top, or by
adding rules to apply_paragraph_rules().
"""

import argparse
import re
import sys
from pathlib import Path

try:
    from docx import Document
    from docx.shared import Pt, RGBColor, Twips, Mm, Cm
    from docx.enum.style import WD_STYLE_TYPE
    from docx.enum.text import WD_BREAK, WD_ALIGN_PARAGRAPH
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn
except ImportError:
    print("python-docx is not installed. Install with: pip install python-docx")
    sys.exit(1)


# Default image paths (resolved relative to the project root)
PHOTOS_DIR = Path(r'C:\Users\mrnig_ndtz4tw\Projects\alpine-map-training\public\photos')
BRAND_DIR  = Path(r'C:\Users\mrnig_ndtz4tw\Projects\alpine-map-training\public\brand')

HERO_DEFAULT    = PHOTOS_DIR / 'lone-skier-navy.jpg'
EMBLEM_DEFAULT  = BRAND_DIR  / 'amt-emblem.png'

# Map "Level N" heading text fragment -> divider image
LEVEL_PHOTO_MAP = {
    'level 1': PHOTOS_DIR / 'skier-beanie-closeup.png',
    'level 2': PHOTOS_DIR / 'pair-with-maps.jpg',
    'level 3': PHOTOS_DIR / 'trio-blue-run.png',
}


# =============================================================================
# TOKENS — Glacier Lab values mapped to Word units
# =============================================================================
# Word colour values are hex strings without '#'.
# Word font sizes are in points (Pt). The script converts internally.
# Spacing is in points; python-docx converts to twips on save.

INK         = '0E1A2E'   # body headings + primary action
INK_2       = '3F4D63'   # body copy on cards
INK_3       = '8693A6'   # meta / captions / mono
INK_4       = 'B5BFCE'   # fill-in lines
RULE        = 'D0D7E0'   # hairline borders
RED         = 'D7263D'   # alpine red
ICE         = '2480B5'   # glacier blue (secondary callouts)
MOSS        = '2E7D5B'   # success
AMBER       = 'C58B2C'   # partial
CRIMSON     = 'A33B2A'   # destructive
PAPER_4     = 'F7F9FB'   # sidebar tint (unused in Word page bg)

FONT_BODY   = 'Manrope'
FONT_MONO   = 'IBM Plex Mono'

# OOXML helpers --------------------------------------------------------------
def _set_qn(el, name, value):
    el.set(qn(name), str(value))


def _add_child(parent, tag, **attrs):
    el = OxmlElement(tag)
    for k, v in attrs.items():
        _set_qn(el, f'w:{k}', v)
    parent.append(el)
    return el


def add_left_border(style, color, size_eighths_pt=24, space_pt=8):
    """Add a left paragraph border to a style (Glacier Lab uses for Aim Body)."""
    pPr = style.element.get_or_add_pPr()
    # Remove any existing pBdr to start clean
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        pPr.remove(existing)
    pBdr = OxmlElement('w:pBdr')
    left = OxmlElement('w:left')
    _set_qn(left, 'w:val', 'single')
    _set_qn(left, 'w:sz', size_eighths_pt)   # eighths of a point: 24 = 3pt
    _set_qn(left, 'w:color', color)
    _set_qn(left, 'w:space', space_pt)
    pBdr.append(left)
    pPr.append(pBdr)


def add_bottom_border(style, color, size_eighths_pt=6, space_pt=6):
    """Add a hairline bottom border — used on page-title (Heading 2)."""
    pPr = style.element.get_or_add_pPr()
    existing = pPr.find(qn('w:pBdr'))
    if existing is not None:
        # only modify the bottom child
        bottom = existing.find(qn('w:bottom'))
        if bottom is not None:
            existing.remove(bottom)
    else:
        existing = OxmlElement('w:pBdr')
        pPr.append(existing)
    bottom = OxmlElement('w:bottom')
    _set_qn(bottom, 'w:val', 'single')
    _set_qn(bottom, 'w:sz', size_eighths_pt)
    _set_qn(bottom, 'w:color', color)
    _set_qn(bottom, 'w:space', space_pt)
    existing.append(bottom)


def set_caps(style):
    """Force uppercase rendering (preserve source text case)."""
    rPr = style.element.get_or_add_rPr()
    if rPr.find(qn('w:caps')) is None:
        _add_child(rPr, 'w:caps')


def set_letter_spacing(style, twentieths_pt):
    """Word spacing in 1/20th of a point. 28 ≈ 1.4pt at 8.5pt body."""
    rPr = style.element.get_or_add_rPr()
    sp = rPr.find(qn('w:spacing'))
    if sp is not None:
        rPr.remove(sp)
    sp = OxmlElement('w:spacing')
    _set_qn(sp, 'w:val', twentieths_pt)
    rPr.append(sp)


def set_font_fallback(style, family, mono=False):
    """Set ascii / hAnsi / cs to the same family. Word falls back per system."""
    rPr = style.element.get_or_add_rPr()
    rFonts = rPr.find(qn('w:rFonts'))
    if rFonts is not None:
        rPr.remove(rFonts)
    rFonts = OxmlElement('w:rFonts')
    _set_qn(rFonts, 'w:ascii', family)
    _set_qn(rFonts, 'w:hAnsi', family)
    _set_qn(rFonts, 'w:cs', family)
    _set_qn(rFonts, 'w:eastAsia', family)
    rPr.append(rFonts)


# =============================================================================
# STYLE DEFINITIONS
# =============================================================================

def _get_or_create_builtin(doc, name, base=None):
    """Return a built-in paragraph style, creating it if the source doc
    didn't define it explicitly (common in cleanly-authored docs)."""
    styles = doc.styles
    try:
        return styles[name]
    except KeyError:
        s = styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
        if base:
            try:
                s.base_style = styles[base]
            except KeyError:
                pass
        return s


def define_styles(doc):
    """Redefine built-in styles + add Glacier Lab custom styles."""
    styles = doc.styles

    # ----- Normal (body) -------------------------------------------------
    normal = _get_or_create_builtin(doc, 'Normal')
    normal.font.name = FONT_BODY
    normal.font.size = Pt(11)
    normal.font.color.rgb = RGBColor.from_string(INK_2)
    normal.paragraph_format.space_before = Pt(0)
    normal.paragraph_format.space_after = Pt(8)
    normal.paragraph_format.line_spacing = 1.45
    set_font_fallback(normal, FONT_BODY)

    # ----- Heading 1 (level opener) -------------------------------------
    h1 = _get_or_create_builtin(doc, 'Heading 1', base='Normal')
    h1.font.name = FONT_BODY
    h1.font.size = Pt(32)
    h1.font.bold = True
    h1.font.color.rgb = RGBColor.from_string(INK)
    h1.paragraph_format.space_before = Pt(40)
    h1.paragraph_format.space_after = Pt(16)
    h1.paragraph_format.line_spacing = 1.05
    set_font_fallback(h1, FONT_BODY)

    # ----- Heading 2 (page title — "B1.1 — Map Purpose…") ----------------
    h2 = _get_or_create_builtin(doc, 'Heading 2', base='Normal')
    h2.font.name = FONT_BODY
    h2.font.size = Pt(22)
    h2.font.bold = True
    h2.font.color.rgb = RGBColor.from_string(INK)
    h2.paragraph_format.space_before = Pt(24)
    h2.paragraph_format.space_after = Pt(10)
    h2.paragraph_format.line_spacing = 1.15
    set_font_fallback(h2, FONT_BODY)
    add_bottom_border(h2, RULE, size_eighths_pt=6, space_pt=8)

    # ----- Heading 3 (section sub-head) ---------------------------------
    h3 = _get_or_create_builtin(doc, 'Heading 3', base='Normal')
    h3.font.name = FONT_BODY
    h3.font.size = Pt(13)
    h3.font.bold = True
    h3.font.color.rgb = RGBColor.from_string(INK)
    h3.paragraph_format.space_before = Pt(18)
    h3.paragraph_format.space_after = Pt(6)
    h3.paragraph_format.line_spacing = 1.2
    set_font_fallback(h3, FONT_BODY)

    # ----- ListParagraph (bullets) --------------------------------------
    lp = _get_or_create_builtin(doc, 'List Paragraph', base='Normal')
    lp.font.name = FONT_BODY
    lp.font.size = Pt(11)
    lp.font.color.rgb = RGBColor.from_string(INK_2)
    lp.paragraph_format.space_before = Pt(2)
    lp.paragraph_format.space_after = Pt(2)
    lp.paragraph_format.line_spacing = 1.45
    set_font_fallback(lp, FONT_BODY)

    # ----- GL Eyebrow ---------------------------------------------------
    gl_eyebrow = _ensure_style(doc, 'GL Eyebrow', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_eyebrow.font.name = FONT_MONO
    gl_eyebrow.font.size = Pt(8.5)
    gl_eyebrow.font.bold = True
    gl_eyebrow.font.color.rgb = RGBColor.from_string(INK_3)
    gl_eyebrow.paragraph_format.space_before = Pt(8)
    gl_eyebrow.paragraph_format.space_after = Pt(4)
    set_font_fallback(gl_eyebrow, FONT_MONO)
    set_caps(gl_eyebrow)
    set_letter_spacing(gl_eyebrow, 28)

    # ----- GL Aim Label -------------------------------------------------
    gl_aim_lbl = _ensure_style(doc, 'GL Aim Label', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_aim_lbl.font.name = FONT_MONO
    gl_aim_lbl.font.size = Pt(8.5)
    gl_aim_lbl.font.bold = True
    gl_aim_lbl.font.color.rgb = RGBColor.from_string(RED)
    gl_aim_lbl.paragraph_format.space_before = Pt(14)
    gl_aim_lbl.paragraph_format.space_after = Pt(3)
    gl_aim_lbl.paragraph_format.left_indent = Pt(12)
    set_font_fallback(gl_aim_lbl, FONT_MONO)
    set_caps(gl_aim_lbl)
    set_letter_spacing(gl_aim_lbl, 28)
    add_left_border(gl_aim_lbl, RED, size_eighths_pt=24, space_pt=8)

    # ----- GL Aim Body --------------------------------------------------
    gl_aim_body = _ensure_style(doc, 'GL Aim Body', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_aim_body.font.name = FONT_BODY
    gl_aim_body.font.size = Pt(12)
    gl_aim_body.font.color.rgb = RGBColor.from_string(INK)
    gl_aim_body.paragraph_format.space_before = Pt(0)
    gl_aim_body.paragraph_format.space_after = Pt(14)
    gl_aim_body.paragraph_format.line_spacing = 1.4
    gl_aim_body.paragraph_format.left_indent = Pt(12)
    set_font_fallback(gl_aim_body, FONT_BODY)
    add_left_border(gl_aim_body, RED, size_eighths_pt=24, space_pt=8)

    # ----- GL Exercise Label --------------------------------------------
    gl_ex = _ensure_style(doc, 'GL Exercise Label', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_ex.font.name = FONT_MONO
    gl_ex.font.size = Pt(8.5)
    gl_ex.font.bold = True
    gl_ex.font.color.rgb = RGBColor.from_string(RED)
    gl_ex.paragraph_format.space_before = Pt(18)
    gl_ex.paragraph_format.space_after = Pt(6)
    set_font_fallback(gl_ex, FONT_MONO)
    set_caps(gl_ex)
    set_letter_spacing(gl_ex, 28)

    # ----- GL Worked Label ----------------------------------------------
    gl_wk = _ensure_style(doc, 'GL Worked Label', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_wk.font.name = FONT_MONO
    gl_wk.font.size = Pt(8.5)
    gl_wk.font.bold = True
    gl_wk.font.color.rgb = RGBColor.from_string(ICE)
    gl_wk.paragraph_format.space_before = Pt(18)
    gl_wk.paragraph_format.space_after = Pt(6)
    set_font_fallback(gl_wk, FONT_MONO)
    set_caps(gl_wk)
    set_letter_spacing(gl_wk, 28)

    # ----- GL Self-Check Label ------------------------------------------
    gl_sc = _ensure_style(doc, 'GL Self-Check Label', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_sc.font.name = FONT_MONO
    gl_sc.font.size = Pt(8.5)
    gl_sc.font.bold = True
    gl_sc.font.color.rgb = RGBColor.from_string(MOSS)
    gl_sc.paragraph_format.space_before = Pt(18)
    gl_sc.paragraph_format.space_after = Pt(6)
    set_font_fallback(gl_sc, FONT_MONO)
    set_caps(gl_sc)
    set_letter_spacing(gl_sc, 28)

    # ----- GL Reflection Label ------------------------------------------
    gl_rf = _ensure_style(doc, 'GL Reflection Label', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_rf.font.name = FONT_MONO
    gl_rf.font.size = Pt(8.5)
    gl_rf.font.bold = True
    gl_rf.font.color.rgb = RGBColor.from_string(INK_3)
    gl_rf.paragraph_format.space_before = Pt(18)
    gl_rf.paragraph_format.space_after = Pt(6)
    set_font_fallback(gl_rf, FONT_MONO)
    set_caps(gl_rf)
    set_letter_spacing(gl_rf, 28)

    # ----- GL Fig Caption -----------------------------------------------
    gl_fc = _ensure_style(doc, 'GL Fig Caption', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_fc.font.name = FONT_MONO
    gl_fc.font.size = Pt(8.5)
    gl_fc.font.bold = True
    gl_fc.font.color.rgb = RGBColor.from_string(INK_3)
    gl_fc.paragraph_format.space_before = Pt(4)
    gl_fc.paragraph_format.space_after = Pt(14)
    set_font_fallback(gl_fc, FONT_MONO)
    set_caps(gl_fc)
    set_letter_spacing(gl_fc, 24)

    # ----- GL Cover Eyebrow ---------------------------------------------
    gl_ce = _ensure_style(doc, 'GL Cover Eyebrow', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_ce.font.name = FONT_MONO
    gl_ce.font.size = Pt(10)
    gl_ce.font.bold = True
    gl_ce.font.color.rgb = RGBColor.from_string(INK_3)
    gl_ce.paragraph_format.space_before = Pt(0)
    gl_ce.paragraph_format.space_after = Pt(8)
    set_font_fallback(gl_ce, FONT_MONO)
    set_caps(gl_ce)
    set_letter_spacing(gl_ce, 40)

    # ----- GL Cover Title -----------------------------------------------
    gl_ct = _ensure_style(doc, 'GL Cover Title', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_ct.font.name = FONT_BODY
    gl_ct.font.size = Pt(64)
    gl_ct.font.bold = True
    gl_ct.font.color.rgb = RGBColor.from_string(INK)
    gl_ct.paragraph_format.space_before = Pt(48)
    gl_ct.paragraph_format.space_after = Pt(20)
    gl_ct.paragraph_format.line_spacing = 1.02
    set_font_fallback(gl_ct, FONT_BODY)

    # ----- GL Cover Sub -------------------------------------------------
    gl_cs = _ensure_style(doc, 'GL Cover Sub', WD_STYLE_TYPE.PARAGRAPH, base='Normal')
    gl_cs.font.name = FONT_BODY
    gl_cs.font.size = Pt(14)
    gl_cs.font.color.rgb = RGBColor.from_string(INK_2)
    gl_cs.paragraph_format.space_before = Pt(0)
    gl_cs.paragraph_format.space_after = Pt(24)
    gl_cs.paragraph_format.line_spacing = 1.4
    set_font_fallback(gl_cs, FONT_BODY)


def _ensure_style(doc, name, style_type, base=None):
    """Get or create a paragraph style."""
    styles = doc.styles
    if name in [s.name for s in styles]:
        return styles[name]
    s = styles.add_style(name, style_type)
    if base:
        s.base_style = styles[base]
    return s


# =============================================================================
# INLINE-OVERRIDE STRIPPING
# =============================================================================
# The original document has inline rPr (run properties) that force Calibri /
# size / colour on every single run. Those win over paragraph styles. Strip
# them so our redefined styles can take effect.

_STRIP_TAGS = ['w:rFonts', 'w:color', 'w:sz', 'w:szCs']


def strip_run_overrides(run, preserve_bold_italic=True):
    rPr = run._element.find(qn('w:rPr'))
    if rPr is None:
        return
    for tag in _STRIP_TAGS:
        for el in rPr.findall(qn(tag)):
            rPr.remove(el)


def strip_paragraph_run_overrides(para):
    for run in para.runs:
        strip_run_overrides(run)


# =============================================================================
# PARAGRAPH RE-TAGGING RULES
# =============================================================================

EYEBROW_RE       = re.compile(r'^\s*LEVEL\s+\d+\s+[·•∙].+[·•∙]\s+PAGE\s+\S+', re.I)
FIG_CAPTION_RE   = re.compile(r'^\s*Fig\.\s+[A-Z0-9.]+', re.I)
EXERCISE_RE      = re.compile(r'^\s*Exercise\s+\d', re.I)
ANSWER_BLANK_RE  = re.compile(r'^_+\s*$|^\s*Answer:\s*_+\s*$')
TICKBOX_RE       = re.compile(r'^\s*\[\s*\]\s+')

# Heading-3 text -> custom paragraph style mapping
H3_TEXT_MAP = {
    'learning aim':     'GL Aim Label',
    'worked example':   'GL Worked Label',
    'self-check':       'GL Self-Check Label',
    'reflection':       'GL Reflection Label',
}


def apply_paragraph_rules(doc):
    """Walk every paragraph and re-tag by content + neighbour context."""
    prev_style_name = None
    for para in doc.paragraphs:
        text = para.text.strip()
        cur_style = para.style.name if para.style else 'Normal'

        new_style = None

        # Never re-tag Heading 1 or Heading 2 — they carry structural meaning
        # (level openers and page titles) and our style redefinitions already
        # restyle them in place.
        if cur_style in ('Heading 1', 'Heading 2'):
            strip_paragraph_run_overrides(para)
            prev_style_name = cur_style
            continue

        # 1. Eyebrow (LEVEL X · ... · PAGE Y) — requires the "PAGE Y" tail
        if EYEBROW_RE.match(text):
            new_style = 'GL Eyebrow'

        # 2. Figure caption (Fig. Lx.x – ...)
        elif FIG_CAPTION_RE.match(text):
            new_style = 'GL Fig Caption'

        # 3. Heading 3 by text
        elif cur_style == 'Heading 3':
            t = text.lower()
            if t in H3_TEXT_MAP:
                new_style = H3_TEXT_MAP[t]
            elif EXERCISE_RE.match(text):
                new_style = 'GL Exercise Label'

        # 4. Aim body (paragraph immediately after Aim Label that's plain body)
        elif prev_style_name == 'GL Aim Label' and cur_style == 'Normal':
            new_style = 'GL Aim Body'

        # 5. Tickbox line
        elif TICKBOX_RE.match(text):
            new_style = 'GL Self-Check'  # only if it exists; otherwise use Normal

        # Apply
        if new_style and new_style in [s.name for s in doc.styles]:
            para.style = doc.styles[new_style]
            prev_style_name = new_style
        else:
            prev_style_name = cur_style

        # Strip inline run overrides so the (re)defined style takes effect
        strip_paragraph_run_overrides(para)


# =============================================================================
# OPTIONAL: COVER PAGE
# =============================================================================

def _move_before(element, target_para):
    """Detach `element` from its current parent and insert it before target_para."""
    parent = element.getparent()
    if parent is not None:
        parent.remove(element)
    target_para._element.addprevious(element)


def _build_text_para(doc, style_name, text=None, runs=None, alignment=None):
    """Create a free-standing <w:p> with given style + text or pre-built runs."""
    new_p = OxmlElement('w:p')
    pPr = OxmlElement('w:pPr')
    pStyle = OxmlElement('w:pStyle')
    _set_qn(pStyle, 'w:val', _style_id(doc, style_name))
    pPr.append(pStyle)
    if alignment is not None:
        jc = OxmlElement('w:jc')
        _set_qn(jc, 'w:val', alignment)
        pPr.append(jc)
    new_p.append(pPr)
    if runs:
        for r in runs:
            new_p.append(r)
    elif text is not None:
        new_p.append(_make_run(text))
    return new_p


def _add_picture_paragraph(doc, image_path, width_mm, alignment='center'):
    """Add a picture to the END of the doc as its own paragraph. Returns the paragraph."""
    para = doc.add_paragraph()
    align_map = {
        'left':   WD_ALIGN_PARAGRAPH.LEFT,
        'center': WD_ALIGN_PARAGRAPH.CENTER,
        'right':  WD_ALIGN_PARAGRAPH.RIGHT,
    }
    para.alignment = align_map.get(alignment, WD_ALIGN_PARAGRAPH.CENTER)
    run = para.add_run()
    run.add_picture(str(image_path), width=Mm(width_mm))
    return para


def _page_break_para():
    """Return a free-standing <w:p> containing only a page break."""
    p = OxmlElement('w:p')
    r = OxmlElement('w:r')
    br = OxmlElement('w:br')
    _set_qn(br, 'w:type', 'page')
    r.append(br)
    p.append(r)
    return p


def insert_cover_page(doc, hero_path=None, emblem_path=None):
    """Insert a styled cover page at the very top of the document.

    Layout (top -> bottom of A4):
      1. Hero photo (full content-width, ~166mm) — if hero_path given and exists
      2. GL Cover Eyebrow line
      3. GL Cover Title ("Read the mountain.")
      4. GL Cover Sub
      5. Spacer + version eyebrow
      6. Small emblem icon (~28mm wide) — if emblem_path given and exists
      7. PerformOS attribution eyebrow
      8. Page break
    """
    first = doc.paragraphs[0] if doc.paragraphs else None
    if first is None:
        return

    # Build cover elements (some at END of doc for images, some free-standing for text),
    # collect their <w:p> elements in display order, then move them all to BEFORE the
    # current first paragraph. The final element in the list is a page break so the
    # cover sits alone on page 1.
    cover_elements = []

    if hero_path and Path(hero_path).exists():
        hero_para = _add_picture_paragraph(doc, hero_path, width_mm=160, alignment='center')
        cover_elements.append(hero_para._element)

    cover_elements.append(_build_text_para(
        doc, 'GL Cover Eyebrow',
        text='ALPINE MAP TRAINING · PERFORMOS · COMPANION MANUAL',
    ))
    cover_elements.append(_build_text_para(
        doc, 'GL Cover Title',
        runs=[_make_run('Read the '), _make_run('mountain.', color=RED)],
    ))
    cover_elements.append(_build_text_para(
        doc, 'GL Cover Sub',
        text='BASI Alpine L4 ISTD navigation — three levels, sixty-plus '
             'workbook pages, one hundred and sixty flashcards, '
             'two graded quizzes.',
    ))
    cover_elements.append(_build_text_para(
        doc, 'GL Cover Eyebrow',
        text='VERSION 11 · MAY 2026',
    ))

    if emblem_path and Path(emblem_path).exists():
        emblem_para = _add_picture_paragraph(doc, emblem_path, width_mm=28, alignment='center')
        cover_elements.append(emblem_para._element)

    cover_elements.append(_build_text_para(
        doc, 'GL Cover Eyebrow',
        text='PERFORMOS · CARTA · GLACIER LAB',
    ))

    # Final element: page break, so the cover sits alone on page 1
    cover_elements.append(_page_break_para())

    # Move each element to just before `first`, preserving display order
    for el in cover_elements:
        _move_before(el, first)


def insert_level_divider_photos(doc, level_photo_map):
    """For each Heading 1 whose text starts with 'Level N', insert a banner
    photo paragraph immediately before it. Photo is matched by case-insensitive
    key lookup in `level_photo_map` against the heading text's prefix.
    """
    # Snapshot current paragraphs first; we'll mutate the doc as we go.
    paragraphs = list(doc.paragraphs)
    inserted_count = 0
    for para in paragraphs:
        if not para.style or para.style.name != 'Heading 1':
            continue
        text = para.text.strip().lower()
        if not text:
            continue
        match_path = None
        for key, path in level_photo_map.items():
            if text.startswith(key):
                match_path = Path(path)
                break
        if match_path is None or not match_path.exists():
            continue

        # Add picture at end of doc, then move just before this Heading 1
        pic_para = _add_picture_paragraph(doc, match_path, width_mm=160, alignment='center')
        _move_before(pic_para._element, para)
        inserted_count += 1

    return inserted_count


def _style_id(doc, name):
    """Return the styleId for a style by display name. python-docx exposes it as .style_id"""
    for s in doc.styles:
        if s.name == name:
            return s.style_id
    return name  # fallback


def _make_run(text, color=None):
    """Build a <w:r> element with text and optional colour override."""
    r = OxmlElement('w:r')
    if color:
        rPr = OxmlElement('w:rPr')
        col = OxmlElement('w:color')
        _set_qn(col, 'w:val', color)
        rPr.append(col)
        r.append(rPr)
    t = OxmlElement('w:t')
    t.text = text
    _set_qn(t, 'xml:space', 'preserve')
    r.append(t)
    return r


# =============================================================================
# OPTIONAL: ANSWER KEY STUB
# =============================================================================

def append_answer_key_stub(doc):
    """Append a starting Answer Key chapter at the back of the document.

    This is a STUB — populate with real answers in a follow-up pass.
    """
    doc.add_paragraph()  # spacer
    # Page break
    p = doc.add_paragraph()
    run = p.add_run()
    run.add_break(WD_BREAK.PAGE)

    p = doc.add_paragraph('ANSWER KEY · ALPINE MAP TRAINING')
    p.style = doc.styles['GL Eyebrow']

    p = doc.add_paragraph('Answer Key')
    p.style = doc.styles['Heading 1']

    p = doc.add_paragraph(
        'Answers are provided here for exercises with determinable correct '
        'responses (matching, multiple-choice, factual recall). Open-ended '
        'exercises are marked as such — discuss in coaching sessions.'
    )

    # Stub for Level 1
    p = doc.add_paragraph('LEVEL 1 · MAP LITERACY · ANSWERS')
    p.style = doc.styles['GL Eyebrow']
    p = doc.add_paragraph('Level 1 — Map literacy')
    p.style = doc.styles['Heading 2']
    p = doc.add_paragraph(
        'TODO: Populate answers for B1.1 through B6.3, the Level 1 Check Quiz, '
        'and the Reflection page.'
    )


# =============================================================================
# MAIN
# =============================================================================

def main():
    parser = argparse.ArgumentParser(description='Re-style workbook to Glacier Lab.')
    parser.add_argument(
        '--input',
        default='Alpine_Map_Training_Companion_Manual_v11.docx',
        help='Source .docx file',
    )
    parser.add_argument(
        '--output',
        default=None,
        help='Output .docx (default: <input>_glacier.docx)',
    )
    parser.add_argument(
        '--cover',
        action='store_true',
        help='Insert a styled cover page at the top',
    )
    parser.add_argument(
        '--hero',
        default=str(HERO_DEFAULT),
        help='Path to the cover hero image (used with --cover). '
             'Set to empty string to disable.',
    )
    parser.add_argument(
        '--emblem',
        default=str(EMBLEM_DEFAULT),
        help='Path to the small brand emblem on the cover (used with --cover). '
             'Set to empty string to disable.',
    )
    parser.add_argument(
        '--level-photos',
        action='store_true',
        help='Insert a banner photo before each Level 1 / Level 2 / Level 3 '
             'Heading 1 (uses the hardcoded LEVEL_PHOTO_MAP).',
    )
    parser.add_argument(
        '--answer-key',
        action='store_true',
        help='Append a stub Answer Key at the back '
             '(skip if your document already has one).',
    )
    args = parser.parse_args()

    in_path = Path(args.input)
    if not in_path.exists():
        print(f'Input not found: {in_path}')
        sys.exit(1)
    out_path = Path(args.output) if args.output else in_path.with_name(in_path.stem + '_glacier.docx')

    print(f'Loading {in_path} …')
    doc = Document(str(in_path))
    print(f'  {len(doc.paragraphs)} paragraphs')

    print('Defining Glacier Lab styles …')
    define_styles(doc)

    print('Re-tagging paragraphs by pattern …')
    apply_paragraph_rules(doc)

    if args.level_photos:
        print('Inserting level-divider photos …')
        n = insert_level_divider_photos(doc, LEVEL_PHOTO_MAP)
        print(f'  Inserted {n} level-divider image(s)')

    if args.cover:
        hero_path = args.hero if args.hero else None
        emblem_path = args.emblem if args.emblem else None
        if hero_path and not Path(hero_path).exists():
            print(f'  warn: hero image not found: {hero_path}')
            hero_path = None
        if emblem_path and not Path(emblem_path).exists():
            print(f'  warn: emblem image not found: {emblem_path}')
            emblem_path = None
        print('Inserting cover page …')
        print(f'  hero:   {hero_path or "(none)"}')
        print(f'  emblem: {emblem_path or "(none)"}')
        insert_cover_page(doc, hero_path=hero_path, emblem_path=emblem_path)

    if args.answer_key:
        print('Appending answer-key stub …')
        append_answer_key_stub(doc)

    print(f'Saving {out_path} …')
    doc.save(str(out_path))
    print('Done.')


if __name__ == '__main__':
    main()
